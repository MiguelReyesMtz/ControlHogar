from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches


RUTA_SALIDA = Path(__file__).resolve().parents[1] / "docs" / "Tecnologias_y_Funcionalidades.docx"


def agregar_bullets(documento, elementos):
    for elemento in elementos:
        documento.add_paragraph(elemento, style="List Bullet")


def configurar_fuente_tabla(tabla):
    for fila in tabla.rows:
        for celda in fila.cells:
            celda.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            for parrafo in celda.paragraphs:
                for run in parrafo.runs:
                    run.font.name = "Arial"


def crear_documento():
    documento = Document()
    estilos = documento.styles
    estilos["Normal"].font.name = "Arial"

    titulo = documento.add_paragraph()
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = titulo.add_run("Control Hogar")
    run.bold = True
    run.font.name = "Arial"

    subtitulo = documento.add_paragraph()
    subtitulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitulo.add_run("Tecnologías usadas, funcionalidades asociadas y flujo de funcionamiento")
    run.bold = True
    run.font.name = "Arial"

    documento.add_paragraph(
        "Este documento describe las tecnologías utilizadas en el sistema, las funcionalidades que habilita cada una "
        "y el flujo general de operación. Las funciones mencionadas corresponden a características implementadas "
        "realmente en el proyecto."
    )

    documento.add_heading("Resumen de tecnologías", level=1)
    tecnologias = [
        ("Node.js", "Ejecuta el servidor de la aplicación y los scripts de simulación IoT."),
        (
            "Express.js",
            "Expone la API REST usada por el frontend para autenticación, dispositivos, usuarios, permisos y administración.",
        ),
        (
            "MongoDB",
            "Almacena de forma persistente hogares, usuarios, dispositivos, permisos y lecturas históricas de sensores.",
        ),
        ("Mongoose", "Define modelos y esquemas para trabajar con MongoDB desde Node.js."),
        ("Socket.IO", "Permite actualizaciones en tiempo real entre servidor y navegador."),
        ("JSON Web Tokens (JWT)", "Mantiene sesiones autenticadas para proteger rutas del backend."),
        ("bcryptjs", "Genera y valida hashes de contraseñas."),
        ("HTML5, CSS3 y JavaScript vanilla", "Construyen la interfaz web sin framework de frontend."),
        ("Canvas API", "Dibuja la gráfica simple de lecturas recientes de temperatura."),
        ("dotenv", "Carga variables de configuración desde el archivo .env."),
        (
            "cors",
            "Configura compatibilidad de solicitudes HTTP entre cliente y servidor cuando sea necesario.",
        ),
    ]

    for nombre, descripcion in tecnologias:
        parrafo = documento.add_paragraph(style="List Bullet")
        parrafo.add_run(nombre + ": ").bold = True
        parrafo.add_run(descripcion)

    documento.add_heading("Funcionalidades asociadas a tecnologías", level=1)

    tabla = documento.add_table(rows=1, cols=3)
    tabla.alignment = WD_TABLE_ALIGNMENT.CENTER
    tabla.style = "Table Grid"
    encabezados = tabla.rows[0].cells
    encabezados[0].text = "Funcionalidad"
    encabezados[1].text = "Tecnología asociada"
    encabezados[2].text = "Cómo participa la tecnología"

    filas = [
        (
            "Registro de usuarios con código de hogar",
            "Express.js, MongoDB, Mongoose, bcryptjs",
            "Express recibe la solicitud de registro; Mongoose consulta el hogar por código; bcryptjs guarda la contraseña como hash; MongoDB persiste el usuario.",
        ),
        (
            "Inicio de sesión",
            "Express.js, MongoDB, Mongoose, bcryptjs, JWT",
            "Express procesa credenciales; Mongoose busca el usuario; bcryptjs compara la contraseña; JWT genera el token de sesión.",
        ),
        (
            "Protección de rutas privadas",
            "JWT, middlewares de Express.js",
            "Los middlewares verifican el token antes de permitir acceso a dispositivos, administración o permisos.",
        ),
        (
            "Creación automática del superadministrador",
            "Node.js, MongoDB, Mongoose, bcryptjs",
            "Al iniciar el servidor, un servicio de semillas crea el hogar y la cuenta admin@hogar.com si no existen.",
        ),
        (
            "Dashboard de dispositivos por usuario",
            "HTML5, CSS3, JavaScript vanilla, API REST",
            "El navegador pide los dispositivos visibles por API REST y JavaScript renderiza las tarjetas del dashboard.",
        ),
        (
            "Control de luces",
            "API REST, Express.js, MongoDB, Socket.IO",
            "El frontend envía un cambio de estado por REST; el backend actualiza MongoDB y emite el cambio por Socket.IO.",
        ),
        (
            "Control de ventilador e intensidad",
            "API REST, Express.js, MongoDB, JavaScript vanilla, Socket.IO",
            "El rango de intensidad se captura en JavaScript, se envía al backend por REST, se guarda en MongoDB y se notifica en tiempo real.",
        ),
        (
            "Lecturas simuladas de temperatura",
            "Node.js, scripts de simulación, MongoDB, Socket.IO",
            "Un simulador ejecutado con setInterval genera variaciones de temperatura, las guarda como lecturas y las publica por Socket.IO.",
        ),
        (
            "Gráfica simple de temperatura",
            "Canvas API, JavaScript vanilla, API REST",
            "JavaScript obtiene lecturas recientes por REST y usa Canvas para dibujar la línea de tendencia.",
        ),
        (
            "Lecturas simuladas de movimiento",
            "Node.js, scripts de simulación, MongoDB, Socket.IO",
            "El simulador genera estados de movimiento detectado o inactivo, los persiste y los envía al cliente en vivo.",
        ),
        (
            "Vista administrativa",
            "HTML5, CSS3, JavaScript vanilla, API REST, JWT",
            "La interfaz muestra secciones administrativas solo a usuarios con rol admin o superadmin, validado también por el backend.",
        ),
        (
            "Agregar dispositivos",
            "API REST, Express.js, MongoDB, Mongoose",
            "El administrador envía nombre, tipo y ubicación; el backend crea el dispositivo en MongoDB mediante Mongoose.",
        ),
        (
            "Eliminar dispositivos",
            "API REST, Express.js, MongoDB, Mongoose, Socket.IO",
            "El backend elimina el dispositivo, sus permisos y lecturas; luego avisa a los clientes conectados por Socket.IO.",
        ),
        (
            "Asignar y revocar permisos",
            "API REST, Express.js, MongoDB, Mongoose, Socket.IO",
            "Los permisos se guardan en MongoDB por usuario y dispositivo; Socket.IO avisa al usuario afectado para refrescar su dashboard.",
        ),
        (
            "Promover o revocar administradores",
            "API REST, Express.js, MongoDB, JWT",
            "El superadministrador cambia el rol del usuario en MongoDB; el backend valida que solo el rol superadmin pueda hacerlo.",
        ),
        (
            "Visualización del código del hogar",
            "API REST, Express.js, MongoDB",
            "La vista administrativa consulta los datos del hogar guardados en MongoDB y muestra el código de registro.",
        ),
        (
            "Configuración por archivo .env",
            "dotenv, Node.js",
            "dotenv carga puerto, URL de MongoDB, secreto JWT, código del hogar e intervalo del simulador.",
        ),
        (
            "Interfaz responsiva",
            "HTML5 y CSS3",
            "CSS organiza formularios, dashboard y panel administrativo para pantallas de escritorio y móviles.",
        ),
    ]

    for funcionalidad, tecnologia, descripcion in filas:
        celdas = tabla.add_row().cells
        celdas[0].text = funcionalidad
        celdas[1].text = tecnologia
        celdas[2].text = descripcion

    for celda in tabla.rows[0].cells:
        for parrafo in celda.paragraphs:
            for run in parrafo.runs:
                run.bold = True

    configurar_fuente_tabla(tabla)

    for fila in tabla.rows:
        fila.cells[0].width = Inches(1.8)
        fila.cells[1].width = Inches(1.9)
        fila.cells[2].width = Inches(3.6)

    documento.add_heading("Flujo de funcionamiento del sistema", level=1)
    documento.add_paragraph(
        "El sistema funciona como una aplicación web cliente-servidor. El navegador muestra la interfaz, "
        "el backend procesa la lógica de negocio, MongoDB conserva los datos y el simulador IoT genera "
        "lecturas periódicas para los sensores."
    )

    flujos = [
        (
            "Arranque del servidor",
            [
                "Node.js ejecuta src/servidor.js.",
                "dotenv carga las variables del archivo .env, incluyendo el puerto, la conexión de MongoDB, el código del hogar y el secreto JWT.",
                "Mongoose abre la conexión con MongoDB.",
                "El servicio de semillas verifica que exista el hogar principal y crea el superadministrador inicial si todavía no existe.",
                "Express registra las rutas REST, sirve los archivos del frontend desde public/ y Socket.IO queda conectado al mismo servidor HTTP.",
                "El simulador IoT inicia un intervalo periódico para generar lecturas de sensores.",
            ],
        ),
        (
            "Registro de un usuario",
            [
                "El usuario llena correo, contraseña y código del hogar desde el frontend.",
                "JavaScript envía la solicitud a la API REST de Express.",
                "El backend valida el correo, la longitud de la contraseña y que el código del hogar exista en MongoDB.",
                "bcryptjs convierte la contraseña en hash antes de guardarla.",
                "Mongoose crea el usuario en MongoDB con rol de usuario común.",
                "jsonwebtoken genera un JWT y el navegador lo guarda localmente para las siguientes solicitudes.",
            ],
        ),
        (
            "Inicio de sesión y autorización",
            [
                "El usuario envía correo y contraseña desde la pantalla de login.",
                "Express busca el usuario en MongoDB usando Mongoose.",
                "bcryptjs compara la contraseña escrita con el hash guardado.",
                "Si las credenciales son correctas, el backend devuelve un JWT.",
                "En cada solicitud protegida, un middleware revisa el JWT y carga el usuario actual.",
                "Las rutas administrativas verifican además que el usuario tenga rol admin o superadmin.",
            ],
        ),
        (
            "Carga del dashboard",
            [
                "El frontend consulta la API REST de dispositivos.",
                "El backend revisa el rol del usuario.",
                "Si es administrador o superadministrador, devuelve todos los dispositivos del hogar.",
                "Si es usuario común, devuelve solo los dispositivos con permiso de visualización.",
                "JavaScript renderiza tarjetas para actuadores y sensores usando HTML y CSS.",
                "Para sensores de temperatura, el frontend consulta lecturas recientes y dibuja una gráfica con Canvas API.",
            ],
        ),
        (
            "Control de actuadores",
            [
                "El usuario presiona el botón de encendido o mueve la intensidad del ventilador.",
                "JavaScript envía el nuevo estado a la API REST.",
                "Express valida que el dispositivo exista y que el usuario tenga permiso de control.",
                "Mongoose actualiza el estado del dispositivo en MongoDB.",
                "Socket.IO emite el cambio a administradores y usuarios autorizados.",
                "Los navegadores conectados actualizan su dashboard sin recargar la página.",
            ],
        ),
        (
            "Simulación de sensores",
            [
                "El simulador IoT se ejecuta dentro del backend con un intervalo configurado en milisegundos.",
                "Para sensores de temperatura, el script genera variaciones graduales dentro de un rango realista.",
                "Para sensores de movimiento, el script alterna entre inactivo y movimiento detectado con cierta probabilidad.",
                "Cada lectura se guarda en MongoDB como histórico de sensor.",
                "Socket.IO envía la nueva lectura a los clientes autorizados.",
                "El dashboard actualiza el valor visible y la gráfica reciente.",
            ],
        ),
        (
            "Administración de dispositivos",
            [
                "Un administrador entra a la vista Administración.",
                "El frontend consulta usuarios, dispositivos y datos del hogar mediante API REST.",
                "Para agregar dispositivos, el administrador indica nombre, tipo y ubicación.",
                "Express crea el dispositivo en MongoDB con una plantilla según su tipo.",
                "Para eliminar dispositivos, el backend borra el dispositivo, sus permisos y sus lecturas históricas.",
                "Socket.IO avisa a los clientes conectados para retirar el dispositivo eliminado de sus vistas.",
            ],
        ),
        (
            "Administración de permisos y roles",
            [
                "El administrador selecciona un usuario común y asigna permisos por dispositivo.",
                "Los permisos se guardan en MongoDB indicando si el usuario puede ver o controlar cada dispositivo.",
                "En sensores solo aplica el permiso de visualización, porque no son actuadores controlables.",
                "Socket.IO notifica al usuario afectado para refrescar su lista de dispositivos.",
                "Solo el superadministrador puede promover usuarios a administradores o revocar ese rol.",
                "El cambio de rol se guarda en MongoDB y afecta las siguientes validaciones de acceso.",
            ],
        ),
    ]

    for titulo_flujo, pasos in flujos:
        documento.add_heading(titulo_flujo, level=2)
        agregar_bullets(documento, pasos)

    documento.add_heading("Alcances reales del sistema", level=1)
    agregar_bullets(
        documento,
        [
            "El sistema usa un solo hogar.",
            "Los dispositivos son simulados; no existe conexión con hardware real.",
            "Los sensores generan datos automáticamente desde scripts del backend.",
            "Los usuarios comunes solo acceden a dispositivos autorizados.",
            "Los administradores tienen acceso completo a los dispositivos.",
            "Solo el superadministrador puede cambiar roles de usuario común a administrador o viceversa.",
        ],
    )

    documento.add_heading("Funciones no implementadas", level=1)
    agregar_bullets(
        documento,
        [
            "Recuperación o cambio de contraseña desde la interfaz.",
            "Edición de dispositivos existentes.",
            "Edición del código del hogar desde la interfaz.",
            "Eliminación de usuarios.",
            "Soporte para varios hogares.",
            "Exportación de lecturas históricas.",
            "Integración con hardware físico real.",
        ],
    )

    RUTA_SALIDA.parent.mkdir(parents=True, exist_ok=True)
    documento.save(RUTA_SALIDA)


if __name__ == "__main__":
    crear_documento()
    print(RUTA_SALIDA)
